from docx import Document
import json
import re
import subprocess
from datetime import datetime


# -------------------------------
# SMART DATE PARSER
# -------------------------------
def parse_date(date_str):

    date_str = date_str.strip()

    formats = [
        "%d.%m.%Y",
        "%d-%m-%Y",
        "%d/%m/%Y",
        "%d.%m.%y",
        "%d-%m-%y",
        "%d/%m/%y",
    ]

    for fmt in formats:
        try:
            return datetime.strptime(date_str, fmt)
        except:
            continue

    nums = re.findall(r"\d+", date_str)

    if len(nums) >= 3:

        day, month, year = nums[0], nums[1], nums[2]

        if len(year) == 2:
            year = "20" + year

        try:
            return datetime(int(year), int(month), int(day))
        except:
            pass

    return None


# -------------------------------
# PARSER
# -------------------------------
def read_weather_doc(file_path):

    doc = Document(file_path)

    result = {
        "location": "",
        "date": "",
        "forecast_period": "",
        "data": [],
    }

    full_text = []

    # -------- PARAGRAPHS --------
    for para in doc.paragraphs:

        lines = para.text.split("\n")

        for line in lines:

            text = line.strip()

            if text:

                full_text.append(text)

                if "WEATHER ADVISORY FOR" in text:
                    result["location"] = text.split("FOR")[-1].strip().title()

                if "Date of Forecast" in text:
                    result["date"] = text.split("–")[-1].strip()

                if "Forecasted period" in text:
                    result["forecast_period"] = text.split(":")[-1].strip()

    # -------- TABLE --------
    table = doc.tables[0]

    table_data = {}

    for row in table.rows:

        key = row.cells[0].text.strip()

        values = [
            cell.text.strip()
            for cell in row.cells[1:]
        ]

        table_data[key] = values

    dates = table_data.get("Date", [])
    rainfall = table_data.get("Rainfall (mm)", [])
    tmax = table_data.get("T-Max (°C)", [])
    tmin = table_data.get("T-Min (°C)", [])
    rhmax = table_data.get("RH Max (%)", [])
    rhmin = table_data.get("RH Min (%)", [])
    wind = table_data.get("Wind Gust (Kph)", [])

    # -------- WARNINGS --------
    warnings_map = {}

    current_day = None

    for line in full_text:

        line = line.replace("–", "-").replace("—", "-")

        match = re.search(
            r"(\d{1,2})\s*(st|nd|rd|th)?\s+[A-Za-z]+\s*-",
            line
        )

        if match:

            current_day = match.group(1).zfill(2)

            warnings_map[current_day] = ""

            continue

        if current_day:

            if line:

                warnings_map[current_day] = line

                current_day = None

    # -------- FINAL BUILD --------
    for i in range(len(dates)):

        dt = parse_date(dates[i])

        if dt:

            day_str = dt.strftime("%d %b")

            day_key = dt.strftime("%d")

        else:

            day_str = dates[i]

            day_key = dates[i][:2]

        result["data"].append({

            "day": day_str,

            "rainfall": float(rainfall[i]),

            "tmax": float(tmax[i]),

            "tmin": float(tmin[i]),

            "rhmax": int(rhmax[i]),

            "rhmin": int(rhmin[i]),

            "wind": wind[i],

            "warning": warnings_map.get(day_key, "")
        })

    return result


# -------------------------------
# RUN
# -------------------------------

file_path = "Barasat_Weather_Advisory095.docx"

data = read_weather_doc(file_path)

# SAVE JSON
with open("weather.json", "w", encoding="utf-8") as f:

    json.dump(data, f, indent=2)

print("✅ weather.json updated")


# -------------------------------
# AUTO GITHUB PUSH
# -------------------------------

subprocess.run("git add .", shell=True)

subprocess.run(
    'git commit -m "daily weather update"',
    shell=True
)

subprocess.run("git push", shell=True)

print("🚀 Uploaded to GitHub successfully")